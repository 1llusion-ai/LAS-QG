import uuid
from pathlib import Path

from src.schemas.types import Document, DocumentType, DocumentStatus


class BaseParser:
    def parse(self, file_path: str) -> Document:
        pass


class TxtParser(BaseParser):
    def parse(self, file_path: str) -> Document:
        path = Path(file_path)
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        return Document(
            id=str(uuid.uuid4()),
            filename=path.name,
            doc_type=DocumentType.TXT,
            content=content,
            status=DocumentStatus.PARSED,
        )


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> Document:
        from docx import Document as DocxDocument

        path = Path(file_path)
        doc = DocxDocument(path)
        content = "\n".join([para.text for para in doc.paragraphs])

        return Document(
            id=str(uuid.uuid4()),
            filename=path.name,
            doc_type=DocumentType.DOCX,
            content=content,
            status=DocumentStatus.PARSED,
        )


class PdfParser(BaseParser):
    def parse(self, file_path: str) -> Document:
        import pdfplumber

        path = Path(file_path)
        content_parts = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content_parts.append(text)

        content = "\n".join(content_parts)

        return Document(
            id=str(uuid.uuid4()),
            filename=path.name,
            doc_type=DocumentType.PDF,
            content=content,
            status=DocumentStatus.PARSED,
        )


class ParserFactory:
    _parsers = {
        DocumentType.TXT: TxtParser(),
        DocumentType.DOCX: DocxParser(),
        DocumentType.PDF: PdfParser(),
    }

    @classmethod
    def get_parser(cls, doc_type: DocumentType) -> BaseParser:
        parser = cls._parsers.get(doc_type)
        if parser is None:
            raise ValueError(f"不支持的文档类型: {doc_type}")
        return parser

    @classmethod
    def parse(cls, file_path: str, doc_type: DocumentType) -> Document:
        parser = cls.get_parser(doc_type)
        return parser.parse(file_path)


def detect_doc_type(filename: str) -> DocumentType:
    ext = Path(filename).suffix.lower()
    type_map = {
        ".txt": DocumentType.TXT,
        ".docx": DocumentType.DOCX,
        ".pdf": DocumentType.PDF,
    }
    doc_type = type_map.get(ext)
    if doc_type is None:
        raise ValueError(f"不支持的文件类型: {ext}")
    return doc_type


def parse_document(file_path: str) -> Document:
    path = Path(file_path)
    doc_type = detect_doc_type(path.name)
    return ParserFactory.parse(file_path, doc_type)
